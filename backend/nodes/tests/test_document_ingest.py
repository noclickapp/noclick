"""
Tests for the shared document-ingest helper (nodes/core/document_ingest.py).

Covers text extraction routing, chunking, deterministic ids, and the
ingest_document / ingest_document_text pipelines (with media resolution,
embedding, and billing mocked).
"""

import io
import types

import pytest
from unittest.mock import Mock, patch

from nodes.core import document_ingest as di


# ---------------------------------------------------------------------------
# Extraction
# ---------------------------------------------------------------------------


def test_extract_plain_text():
    assert di.extract_text(b"hello world", "text/plain", "a.txt") == "hello world"


def test_extract_json_and_csv_decoded():
    assert "key" in di.extract_text(b'{"key": 1}', "application/json", "a.json")
    assert "a,b" in di.extract_text(b"a,b\n1,2", "text/csv", "a.csv")


def test_extract_html_strips_tags():
    out = di.extract_text(b"<html><body><p>Hi</p><p>There</p></body></html>", "text/html", "a.html")
    assert "Hi" in out and "There" in out
    assert "<p>" not in out


def test_extract_docx_roundtrip():
    import docx

    document = docx.Document()
    document.add_paragraph("First line")
    document.add_paragraph("Second line")
    buf = io.BytesIO()
    document.save(buf)
    out = di.extract_text(buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "a.docx")
    assert "First line" in out and "Second line" in out


def test_extract_unsupported_raises():
    with pytest.raises(ValueError, match="Unsupported document type"):
        di.extract_text(b"\x00\x01binary", "application/octet-stream", "a.bin")


def test_document_accept_matches_supported_types():
    accepted = set(di.DOCUMENT_ACCEPT.split(","))
    assert ".doc" not in accepted
    assert ".docx" in accepted
    assert ".pdf" in accepted


def test_pdf_routes_to_pypdf():
    # Avoid generating a real PDF; assert routing calls pypdf with a fake reader.
    fake_page = types.SimpleNamespace(extract_text=lambda: "page text")
    fake_reader = types.SimpleNamespace(pages=[fake_page, fake_page])
    with patch("pypdf.PdfReader", return_value=fake_reader):
        out = di.extract_text(b"%PDF-1.4 ...", "application/pdf", "a.pdf")
    assert out == "page text\n\npage text"


# ---------------------------------------------------------------------------
# Chunking + ids
# ---------------------------------------------------------------------------


def test_chunk_short_text_single_chunk():
    assert di.chunk_text("short", 1000, 200) == ["short"]


def test_chunk_empty():
    assert di.chunk_text("", 1000, 200) == []
    assert di.chunk_text("   ", 1000, 200) == []


def test_chunk_long_text_overlaps():
    text = "word. " * 500  # 3000 chars
    chunks = di.chunk_text(text, 1000, 200)
    assert len(chunks) >= 3
    assert all(len(c) <= 1000 for c in chunks)


def test_chunk_id_deterministic():
    a = di.chunk_id("doc.txt", 0)
    b = di.chunk_id("doc.txt", 0)
    c = di.chunk_id("doc.txt", 1)
    assert a == b and a != c
    import uuid

    uuid.UUID(a)  # valid uuid


# ---------------------------------------------------------------------------
# Pipelines
# ---------------------------------------------------------------------------


def _fake_resolved(data=b"hello world. " * 200, mime="text/plain", filename="doc.txt"):
    return types.SimpleNamespace(data=data, mime_type=mime, filename=filename, download_url=None)


async def _fake_aembedding(model, input):
    return types.SimpleNamespace(
        data=[{"embedding": [0.1, 0.2, 0.3]} for _ in input],
        usage=types.SimpleNamespace(total_tokens=len(input) * 5),
    )


class _FakeNode:
    user_id = "11111111-1111-1111-1111-111111111111"
    organization_id = None
    sio = None
    sid = None
    node_id = "n1"
    workflow_id = "w1"


@pytest.mark.asyncio
async def test_ingest_document_embeds_and_charges():
    tracker = Mock()
    with patch("nodes.core.media_resolver.resolve_media_input", return_value=_fake_resolved()), \
         patch("litellm.aembedding", side_effect=_fake_aembedding), \
         patch("billing.usage_tracker.usage_tracker", tracker):
        records, info = await di.ingest_document(_FakeNode(), "res-id", chunk_size=1000, chunk_overlap=200)

    assert len(records) >= 2
    r = records[0]
    assert set(r.keys()) == {"id", "values", "metadata"}
    assert r["values"] == [0.1, 0.2, 0.3]
    assert r["metadata"]["source"] == "doc.txt"
    assert r["metadata"]["chunk_index"] == 0
    assert "text" in r["metadata"]
    # Gate + charge both fired.
    assert tracker.enforce_credit_gate.called
    assert tracker.track_usage_event.called


@pytest.mark.asyncio
async def test_ingest_document_text_no_embedding_no_charge():
    tracker = Mock()
    with patch("nodes.core.media_resolver.resolve_media_input", return_value=_fake_resolved()), \
         patch("billing.usage_tracker.usage_tracker", tracker):
        records, info = await di.ingest_document_text("res-id", chunk_size=1000, chunk_overlap=200)

    assert len(records) >= 2
    r = records[0]
    assert set(r.keys()) == {"id", "text", "metadata"}
    assert r["metadata"]["source"] == "doc.txt"
    # Server-side embedding path never touches NoClick billing.
    assert not tracker.track_usage_event.called


@pytest.mark.asyncio
async def test_ingest_empty_document_returns_no_records():
    with patch("nodes.core.media_resolver.resolve_media_input", return_value=_fake_resolved(data=b"   ")):
        records, info = await di.ingest_document_text("res-id")
    assert records == []


def test_embedding_cost_defaults_to_zero_without_operator_policy():
    from billing.pricing import get_embedding_cost
    from decimal import Decimal

    # Community defaults record no commercial price; an operator policy may
    # register its own accounting implementation.
    assert get_embedding_cost(1_000_000) == Decimal("0")
    assert get_embedding_cost(0) == Decimal("0")
